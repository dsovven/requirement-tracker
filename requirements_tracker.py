"""
Requirements Tracker - PDF Requirement Capture Tool

Opens PDF files, allows rectangular screenshot capture of requirements,
stamps requirement numbers on the PDF, and generates a tracking document.
"""

import sys
import os
import json
import base64
from datetime import datetime
from io import BytesIO
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

import fitz  # PyMuPDF

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QGraphicsView, QGraphicsScene, QToolBar, QAction, QFileDialog,
    QLabel, QLineEdit, QCheckBox, QListWidget, QListWidgetItem,
    QSplitter, QMessageBox, QPushButton, QScrollArea, QSizePolicy,
    QStyle, QGroupBox, QFormLayout, QDialog, QSlider, QButtonGroup,
    QMenu, QInputDialog, QTextEdit
)
from PyQt5.QtCore import (
    Qt, QRectF, QRect, QPoint, QSize, pyqtSignal, QBuffer, QByteArray
)
from PyQt5.QtGui import (
    QPixmap, QImage, QPainter, QPen, QColor, QBrush, QFont,
    QKeySequence, QIcon
)

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pytesseract
    from PIL import Image
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

try:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XlImage
    from openpyxl.styles import Font as XlFont, Alignment as XlAlignment
    from openpyxl.utils.units import pixels_to_EMU
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


# ---------------------------------------------------------------------------
# Data
# ---------------------------------------------------------------------------

# Preset colors for markup stamp and outline (name, RGB 0-1 tuple, hex for UI)
MARKUP_COLORS = [
    ("Red",     (0.85, 0.15, 0.15), "#d92626"),
    ("Blue",    (0.15, 0.30, 0.85), "#264dd9"),
    ("Green",   (0.10, 0.55, 0.20), "#1a8c33"),
    ("Purple",  (0.55, 0.15, 0.70), "#8c26b3"),
    ("Orange",  (0.90, 0.45, 0.05), "#e6730d"),
]


@dataclass
class Requirement:
    number: str          # "1", "2", "7.1", etc.
    screenshot: QPixmap  # cropped image of the requirement area
    page: int            # 0-based page index
    pdf_rect: tuple      # (x0, y0, x1, y1) in PDF points
    text: str = ""       # extracted text from the captured region
    edited: bool = False # True if screenshot was modified in the editor
    highlight_pixmap: Optional[QPixmap] = None  # screenshot with highlights only (no white-out)
    markup_color: tuple = (0.85, 0.15, 0.15)  # RGB for stamp/outline on PDF
    notes: str = ""      # user-entered notes (exported to Notes column)


def pixmap_to_bytes(pixmap: QPixmap) -> BytesIO:
    """Convert a QPixmap to a BytesIO PNG stream."""
    ba = QByteArray()
    buf = QBuffer(ba)
    buf.open(QBuffer.WriteOnly)
    pixmap.save(buf, "PNG")
    buf.close()
    bio = BytesIO(ba.data())
    bio.seek(0)
    return bio


# ---------------------------------------------------------------------------
# PDF Page Widget  (handles rendering + rectangle drawing)
# ---------------------------------------------------------------------------

class PDFPageWidget(QWidget):
    """Displays a single PDF page pixmap and lets the user draw rectangles."""

    selection_made = pyqtSignal(QRectF)   # rectangle in *pixmap* coordinates
    zoom_requested = pyqtSignal(int)      # wheel delta for zoom

    def __init__(self, parent=None):
        super().__init__(parent)
        self._pixmap: Optional[QPixmap] = None
        self._drawing = False
        self._start: Optional[QPoint] = None
        self._current_rect: Optional[QRect] = None
        self.setCursor(Qt.CrossCursor)

    def set_pixmap(self, pixmap: QPixmap):
        self._pixmap = pixmap
        self.setFixedSize(pixmap.size())
        self.update()

    # -- painting ----------------------------------------------------------

    def paintEvent(self, event):
        if not self._pixmap:
            return
        painter = QPainter(self)
        painter.drawPixmap(0, 0, self._pixmap)
        if self._current_rect:
            pen = QPen(QColor(137, 180, 250), 2, Qt.DashLine)
            painter.setPen(pen)
            painter.setBrush(QBrush(QColor(137, 180, 250, 30)))
            painter.drawRect(self._current_rect)

    # -- mouse events ------------------------------------------------------

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._drawing = True
            self._start = event.pos()
            self._current_rect = QRect(self._start, self._start)

    def mouseMoveEvent(self, event):
        if self._drawing:
            self._current_rect = QRect(self._start, event.pos()).normalized()
            self.update()

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.LeftButton and self._drawing:
            self._drawing = False
            rect = QRect(self._start, event.pos()).normalized()
            self._current_rect = None
            self.update()
            if rect.width() > 10 and rect.height() > 10:
                self.selection_made.emit(QRectF(rect))

    def wheelEvent(self, event):
        if event.modifiers() & Qt.ControlModifier:
            self.zoom_requested.emit(event.angleDelta().y())
            event.accept()
        else:
            event.ignore()  # propagate to scroll area


# ---------------------------------------------------------------------------
# Zoomable Scroll Area
# ---------------------------------------------------------------------------

class ZoomScrollArea(QScrollArea):
    """QScrollArea that forwards Ctrl+Wheel to a zoom signal."""
    zoom_requested = pyqtSignal(int)

    def wheelEvent(self, event):
        if event.modifiers() & Qt.ControlModifier:
            self.zoom_requested.emit(event.angleDelta().y())
            event.accept()
        else:
            super().wheelEvent(event)


# ---------------------------------------------------------------------------
# Screenshot Editor Dialog  (highlight / white-out annotation)
# ---------------------------------------------------------------------------

class _EditorCanvas(QLabel):
    """Label that supports freehand and rectangle drawing on its pixmap."""

    HIGHLIGHT_ALPHA = 0.30  # opacity applied uniformly to highlight strokes

    def __init__(self, pixmap: QPixmap, parent=None):
        super().__init__(parent)
        self._pixmap = pixmap.copy()
        self._original = pixmap.copy()  # pristine copy for highlight-only export
        self.setPixmap(self._pixmap)
        self.setFixedSize(pixmap.size())
        self.setCursor(Qt.CrossCursor)

        self._drawing = False
        self._last_point: Optional[QPoint] = None
        self._start_point: Optional[QPoint] = None
        self._undo_stack: list = []  # list[QPixmap]
        self._max_undo = 20

        # cumulative highlight-only layer (transparent, only highlight strokes)
        self._highlight_layer = QPixmap(pixmap.size())
        self._highlight_layer.fill(Qt.transparent)

        # stroke overlay for highlight (prevents alpha accumulation)
        self._stroke_base: Optional[QPixmap] = None
        self._stroke_overlay: Optional[QPixmap] = None

        # tool settings (set by parent dialog)
        self.brush_size = 20
        self.tool = "highlight"  # "highlight" or "whiteout"
        self.draw_mode = "rectangle"  # "brush" or "rectangle"
        self.highlight_color = QColor(255, 255, 0)  # yellow default

    def get_pixmap(self) -> QPixmap:
        return self._pixmap.copy()

    def get_highlight_pixmap(self) -> QPixmap:
        """Return the original screenshot with only highlights composited."""
        result = self._original.copy()
        painter = QPainter(result)
        painter.setOpacity(self.HIGHLIGHT_ALPHA)
        painter.drawPixmap(0, 0, self._highlight_layer)
        painter.end()
        return result

    def undo(self):
        if self._undo_stack:
            self._pixmap = self._undo_stack.pop()
            self.setPixmap(self._pixmap)

    def _push_undo(self):
        if len(self._undo_stack) >= self._max_undo:
            self._undo_stack.pop(0)
        self._undo_stack.append(self._pixmap.copy())

    # -- overlay helpers for highlight tool --------------------------------

    def _begin_highlight_stroke(self):
        """Snapshot base and create a transparent overlay for the stroke."""
        self._stroke_base = self._pixmap.copy()
        self._stroke_overlay = QPixmap(self._pixmap.size())
        self._stroke_overlay.fill(Qt.transparent)

    def _draw_highlight_on_overlay(self, p1: QPoint, p2: QPoint = None):
        """Draw highlight color onto the overlay (no alpha — alpha applied at composite)."""
        painter = QPainter(self._stroke_overlay)
        painter.setRenderHint(QPainter.Antialiasing)
        painter.setPen(Qt.NoPen)
        painter.setBrush(self.highlight_color)
        if p2 is None:
            r = self.brush_size / 2
            painter.drawEllipse(p1, r, r)
        else:
            pen = QPen(self.highlight_color, self.brush_size,
                       Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin)
            painter.setPen(pen)
            painter.drawLine(p1, p2)
        painter.end()
        self._composite_highlight()

    def _draw_highlight_rect_on_overlay(self, rect: QRect):
        """Draw a filled highlight rectangle onto the overlay."""
        self._stroke_overlay.fill(Qt.transparent)
        painter = QPainter(self._stroke_overlay)
        painter.setPen(Qt.NoPen)
        painter.setBrush(self.highlight_color)
        painter.drawRect(rect)
        painter.end()
        self._composite_highlight()

    def _composite_highlight(self):
        """Blend overlay onto base at uniform alpha and display."""
        self._pixmap = self._stroke_base.copy()
        painter = QPainter(self._pixmap)
        painter.setOpacity(self.HIGHLIGHT_ALPHA)
        painter.drawPixmap(0, 0, self._stroke_overlay)
        painter.end()
        self.setPixmap(self._pixmap)

    def _finish_highlight_stroke(self):
        """Bake the composited result and accumulate onto highlight layer."""
        if self._stroke_overlay is not None:
            painter = QPainter(self._highlight_layer)
            painter.drawPixmap(0, 0, self._stroke_overlay)
            painter.end()
        self._stroke_base = None
        self._stroke_overlay = None

    # -- direct drawing for white-out tool --------------------------------

    def _draw_whiteout(self, p1: QPoint, p2: QPoint = None):
        painter = QPainter(self._pixmap)
        painter.setRenderHint(QPainter.Antialiasing)
        painter.setCompositionMode(QPainter.CompositionMode_Source)
        if p2 is None:
            painter.setPen(Qt.NoPen)
            painter.setBrush(QColor(255, 255, 255))
            r = self.brush_size / 2
            painter.drawEllipse(p1, r, r)
        else:
            pen = QPen(QColor(255, 255, 255), self.brush_size,
                       Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin)
            painter.setPen(pen)
            painter.drawLine(p1, p2)
        painter.end()
        self.setPixmap(self._pixmap)

    def _draw_whiteout_rect(self, rect: QRect):
        """Fill a rectangle with solid white."""
        painter = QPainter(self._pixmap)
        painter.setPen(Qt.NoPen)
        painter.setBrush(QColor(255, 255, 255))
        painter.drawRect(rect)
        painter.end()
        self.setPixmap(self._pixmap)

    def _preview_whiteout_rect(self, rect: QRect):
        """Show a preview of the white-out rectangle without baking it."""
        preview = self._stroke_base.copy()
        painter = QPainter(preview)
        painter.setPen(QPen(QColor(180, 180, 180), 1, Qt.DashLine))
        painter.setBrush(QColor(255, 255, 255, 180))
        painter.drawRect(rect)
        painter.end()
        self._pixmap = preview
        self.setPixmap(self._pixmap)

    # -- mouse events -----------------------------------------------------

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self._push_undo()
            self._drawing = True
            self._start_point = event.pos()
            self._last_point = event.pos()

            if self.draw_mode == "brush":
                if self.tool == "highlight":
                    self._begin_highlight_stroke()
                    self._draw_highlight_on_overlay(event.pos())
                else:
                    self._draw_whiteout(event.pos())
            else:
                # rectangle mode — snapshot base for live preview
                if self.tool == "highlight":
                    self._begin_highlight_stroke()
                else:
                    self._stroke_base = self._pixmap.copy()

    def mouseMoveEvent(self, event):
        if self._drawing and self._start_point:
            if self.draw_mode == "brush":
                if self.tool == "highlight":
                    self._draw_highlight_on_overlay(self._last_point, event.pos())
                else:
                    self._draw_whiteout(self._last_point, event.pos())
                self._last_point = event.pos()
            else:
                # rectangle mode — live preview
                rect = QRect(self._start_point, event.pos()).normalized()
                if self.tool == "highlight":
                    self._draw_highlight_rect_on_overlay(rect)
                else:
                    self._preview_whiteout_rect(rect)

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.LeftButton and self._drawing:
            self._drawing = False

            if self.draw_mode == "rectangle" and self._start_point:
                rect = QRect(self._start_point, event.pos()).normalized()
                if rect.width() > 2 and rect.height() > 2:
                    if self.tool == "highlight":
                        self._finish_highlight_stroke()
                    else:
                        # bake the final rectangle from the clean base
                        self._pixmap = self._stroke_base.copy()
                        self._draw_whiteout_rect(rect)
                        self._stroke_base = None
                else:
                    # too small — revert
                    if self.tool == "highlight":
                        self._pixmap = self._stroke_base.copy()
                        self.setPixmap(self._pixmap)
                        self._stroke_base = None
                        self._stroke_overlay = None
                    elif self._stroke_base:
                        self._pixmap = self._stroke_base.copy()
                        self.setPixmap(self._pixmap)
                        self._stroke_base = None
                    if self._undo_stack:
                        self._undo_stack.pop()
            else:
                if self.tool == "highlight":
                    self._finish_highlight_stroke()

            self._last_point = None
            self._start_point = None


class ScreenshotEditorDialog(QDialog):
    """Dialog for annotating a screenshot with highlight or white-out."""

    def __init__(self, pixmap: QPixmap, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Edit Screenshot")
        self.setMinimumSize(600, 400)
        self.resize(
            min(pixmap.width() + 60, 1200),
            min(pixmap.height() + 100, 800),
        )

        layout = QVBoxLayout(self)

        # -- toolbar row --
        toolbar = QHBoxLayout()

        self._btn_highlight = QPushButton("Highlight")
        self._btn_highlight.setCheckable(True)
        self._btn_highlight.setChecked(True)
        self._btn_highlight.setStyleSheet(
            "QPushButton:checked { background: #854d0e; color: #fef08a; "
            "font-weight: bold; border-color: #fef08a; }"
        )

        self._btn_whiteout = QPushButton("White-out")
        self._btn_whiteout.setCheckable(True)
        self._btn_whiteout.setStyleSheet(
            "QPushButton:checked { background: #45475a; color: #cdd6f4; "
            "font-weight: bold; border-color: #cdd6f4; }"
        )

        self._tool_group = QButtonGroup(self)
        self._tool_group.setExclusive(True)
        self._tool_group.addButton(self._btn_highlight, 0)
        self._tool_group.addButton(self._btn_whiteout, 1)

        toolbar.addWidget(self._btn_highlight)
        toolbar.addWidget(self._btn_whiteout)

        toolbar.addSpacing(16)

        # highlight color buttons
        self._highlight_colors = [
            ("Yellow", QColor(255, 255, 0), "#854d0e", "#fef08a"),
            ("Orange", QColor(255, 165, 0), "#7c2d12", "#fed7aa"),
            ("Green", QColor(0, 255, 0), "#14532d", "#bbf7d0"),
        ]
        self._color_group = QButtonGroup(self)
        self._color_group.setExclusive(True)
        for i, (name, _qc, bg_dark, fg) in enumerate(self._highlight_colors):
            btn = QPushButton(name)
            btn.setCheckable(True)
            btn.setStyleSheet(
                f"QPushButton:checked {{ background: {bg_dark}; color: {fg}; "
                f"font-weight: bold; border-color: {fg}; }}"
            )
            if i == 0:
                btn.setChecked(True)
            self._color_group.addButton(btn, i)
            toolbar.addWidget(btn)

        toolbar.addSpacing(16)

        # draw mode toggle
        self._btn_brush = QPushButton("Brush")
        self._btn_brush.setCheckable(True)
        self._btn_brush.setStyleSheet(
            "QPushButton:checked { background: #1e3a5f; color: #89b4fa; "
            "font-weight: bold; border-color: #89b4fa; }"
        )

        self._btn_rect = QPushButton("Rectangle")
        self._btn_rect.setCheckable(True)
        self._btn_rect.setChecked(True)
        self._btn_rect.setStyleSheet(
            "QPushButton:checked { background: #1e3a5f; color: #89b4fa; "
            "font-weight: bold; border-color: #89b4fa; }"
        )

        self._mode_group = QButtonGroup(self)
        self._mode_group.setExclusive(True)
        self._mode_group.addButton(self._btn_brush, 0)
        self._mode_group.addButton(self._btn_rect, 1)

        toolbar.addWidget(self._btn_brush)
        toolbar.addWidget(self._btn_rect)

        toolbar.addSpacing(16)

        self._brush_size_label = QLabel("Brush:")
        toolbar.addWidget(self._brush_size_label)
        self._size_slider = QSlider(Qt.Horizontal)
        self._size_slider.setRange(5, 50)
        self._size_slider.setValue(20)
        self._size_slider.setFixedWidth(120)
        toolbar.addWidget(self._size_slider)

        self._size_label = QLabel("20px")
        self._size_label.setFixedWidth(40)
        toolbar.addWidget(self._size_label)

        toolbar.addSpacing(16)

        self._btn_undo = QPushButton("Undo")
        toolbar.addWidget(self._btn_undo)

        toolbar.addStretch()

        self._btn_save = QPushButton("Save")
        self._btn_save.setStyleSheet(
            "QPushButton { background: #89b4fa; color: #11111b; font-weight: bold; "
            "padding: 6px 20px; border: none; border-radius: 6px; }"
            "QPushButton:hover { background: #b4d0fb; }"
            "QPushButton:pressed { background: #74a8f7; }"
        )
        self._btn_cancel = QPushButton("Cancel")

        toolbar.addWidget(self._btn_save)
        toolbar.addWidget(self._btn_cancel)

        layout.addLayout(toolbar)

        # -- canvas in scroll area --
        self._canvas = _EditorCanvas(pixmap)
        scroll = QScrollArea()
        scroll.setWidget(self._canvas)
        scroll.setAlignment(Qt.AlignCenter)
        layout.addWidget(scroll, 1)

        # -- connections --
        self._tool_group.buttonClicked[int].connect(self._on_tool_changed)
        self._color_group.buttonClicked[int].connect(self._on_color_changed)
        self._mode_group.buttonClicked[int].connect(self._on_mode_changed)
        self._size_slider.valueChanged.connect(self._on_size_changed)
        self._btn_undo.clicked.connect(self._canvas.undo)
        self._btn_save.clicked.connect(self.accept)
        self._btn_cancel.clicked.connect(self.reject)

        # apply initial defaults (rectangle mode — disable brush size)
        self._size_slider.setEnabled(False)
        self._brush_size_label.setEnabled(False)
        self._size_label.setEnabled(False)

    def _on_tool_changed(self, btn_id):
        self._canvas.tool = "highlight" if btn_id == 0 else "whiteout"

    def _on_color_changed(self, btn_id):
        _name, qc, _bg, _fg = self._highlight_colors[btn_id]
        self._canvas.highlight_color = qc

    def _on_mode_changed(self, btn_id):
        self._canvas.draw_mode = "brush" if btn_id == 0 else "rectangle"
        is_brush = btn_id == 0
        self._size_slider.setEnabled(is_brush)
        self._brush_size_label.setEnabled(is_brush)
        self._size_label.setEnabled(is_brush)

    def _on_size_changed(self, val):
        self._canvas.brush_size = val
        self._size_label.setText(f"{val}px")

    def get_pixmap(self) -> QPixmap:
        return self._canvas.get_pixmap()

    def get_highlight_pixmap(self) -> QPixmap:
        """Return original screenshot with only highlights applied (no white-out)."""
        return self._canvas.get_highlight_pixmap()


# ---------------------------------------------------------------------------
# PDF Viewer  (scroll area + page widget + navigation helpers)
# ---------------------------------------------------------------------------

class PDFViewer(QWidget):
    """Composite widget: scroll area with a PDFPageWidget inside."""

    selection_made = pyqtSignal(int, QRectF)  # (page, rect in pixmap coords)
    page_changed = pyqtSignal(int, int)       # (current 0-based, total)

    RENDER_ZOOM = 2.0   # default render resolution multiplier

    def __init__(self, parent=None):
        super().__init__(parent)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        self._scroll = ZoomScrollArea()
        self._scroll.setWidgetResizable(False)
        self._scroll.setAlignment(Qt.AlignCenter)
        self._page_widget = PDFPageWidget()
        self._scroll.setWidget(self._page_widget)
        layout.addWidget(self._scroll)

        # internal state
        self._doc: Optional[fitz.Document] = None
        self._current_page = 0
        self._zoom = self.RENDER_ZOOM

        # signals
        self._page_widget.selection_made.connect(self._on_selection)
        self._page_widget.zoom_requested.connect(self._on_zoom_wheel)
        self._scroll.zoom_requested.connect(self._on_zoom_wheel)

    # -- public API --------------------------------------------------------

    @property
    def render_zoom(self):
        return self._zoom

    @property
    def current_page(self):
        return self._current_page

    @property
    def total_pages(self):
        return len(self._doc) if self._doc else 0

    def set_document(self, doc: fitz.Document, page: int = None):
        self._doc = doc
        if page is not None:
            self._current_page = page
        self._current_page = max(0, min(self._current_page, self.total_pages - 1))
        self._render()

    def next_page(self):
        if self._current_page < self.total_pages - 1:
            self._current_page += 1
            self._render()

    def prev_page(self):
        if self._current_page > 0:
            self._current_page -= 1
            self._render()

    def go_to_page(self, page: int):
        page = max(0, min(page, self.total_pages - 1))
        if page != self._current_page:
            self._current_page = page
            self._render()

    def zoom_in(self):
        self._zoom = min(self._zoom * 1.25, 8.0)
        self._render()

    def zoom_out(self):
        self._zoom = max(self._zoom / 1.25, 0.5)
        self._render()

    def fit_width(self):
        if not self._doc:
            return
        page = self._doc[self._current_page]
        vp_w = self._scroll.viewport().width() - 20  # small margin
        self._zoom = max(vp_w / page.rect.width, 0.5)
        self._render()

    @property
    def page_rect_origin(self):
        """Return the (x0, y0) of the current page's rect (CropBox/rotation offset)."""
        if self._doc and 0 <= self._current_page < len(self._doc):
            r = self._doc[self._current_page].rect
            return (r.x0, r.y0)
        return (0.0, 0.0)

    def scroll_to_pdf_point(self, pdf_x, pdf_y):
        ox, oy = self.page_rect_origin
        px = int((pdf_x - ox) * self._zoom)
        py = int((pdf_y - oy) * self._zoom)
        self._scroll.ensureVisible(px, py, 100, 100)

    # -- internal ----------------------------------------------------------

    def _render(self):
        if not self._doc:
            return
        page = self._doc[self._current_page]
        mat = fitz.Matrix(self._zoom, self._zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = QImage(pix.samples, pix.width, pix.height,
                     pix.stride, QImage.Format_RGB888).copy()
        self._page_widget.set_pixmap(QPixmap.fromImage(img))
        self.page_changed.emit(self._current_page, self.total_pages)

    def _on_selection(self, rect: QRectF):
        self.selection_made.emit(self._current_page, rect)

    def _on_zoom_wheel(self, delta):
        if delta > 0:
            self.zoom_in()
        else:
            self.zoom_out()


# ---------------------------------------------------------------------------
# Requirement List Item Widget
# ---------------------------------------------------------------------------

class ReqItemWidget(QWidget):
    """Custom widget shown for each requirement in the list."""

    def __init__(self, req: Requirement, parent=None):
        super().__init__(parent)
        self.setStyleSheet("background: transparent;")
        layout = QHBoxLayout(self)
        layout.setContentsMargins(8, 6, 8, 6)
        layout.setSpacing(10)

        # number badge
        num = QLabel(req.number)
        num.setFont(QFont("Segoe UI", 12, QFont.Bold))
        r, g, b = req.markup_color
        hexc = f"#{int(r*255):02x}{int(g*255):02x}{int(b*255):02x}"
        num.setStyleSheet(
            f"color: {hexc}; background: transparent;"
        )
        num.setFixedWidth(48)
        num.setAlignment(Qt.AlignCenter)

        # thumbnail with rounded border
        thumb = QLabel()
        scaled = req.screenshot.scaled(
            120, 80, Qt.KeepAspectRatio, Qt.SmoothTransformation
        )
        thumb.setPixmap(scaled)
        thumb.setFixedSize(120, 80)
        thumb.setAlignment(Qt.AlignCenter)
        thumb.setStyleSheet(
            "border: 1px solid #45475a; border-radius: 4px; "
            "background: #11111b;"
        )

        # right column: page info + notes
        right_col = QVBoxLayout()
        right_col.setSpacing(2)

        info = QLabel(f"Page {req.page + 1}")
        info.setFont(QFont("Segoe UI", 9))
        info.setStyleSheet("color: #a6adc8; background: transparent;")
        info.setAlignment(Qt.AlignCenter)
        right_col.addWidget(info)

        if req.notes:
            notes_lbl = QLabel("notes")
            notes_lbl.setFont(QFont("Segoe UI", 8))
            notes_lbl.setStyleSheet(
                "color: #89b4fa; background: transparent;"
            )
            notes_lbl.setToolTip(req.notes[:200])
            notes_lbl.setAlignment(Qt.AlignCenter)
            right_col.addWidget(notes_lbl)

        right_col.addStretch()

        layout.addWidget(num)
        layout.addWidget(thumb)
        layout.addLayout(right_col)


# ---------------------------------------------------------------------------
# Requirements Panel  (right-hand sidebar)
# ---------------------------------------------------------------------------

class RequirementsPanel(QWidget):
    """Sidebar listing captured requirements + numbering controls."""

    delete_requested = pyqtSignal(int)  # list row index
    bulk_delete_requested = pyqtSignal(list)  # list of row indices
    edit_requested = pyqtSignal(int)    # list row index (double-click)
    edit_notes_requested = pyqtSignal(int)
    change_color_requested = pyqtSignal(int, int)  # row, color index
    copy_text_requested = pyqtSignal(int)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setStyleSheet("RequirementsPanel { background: #1e1e2e; }")
        layout = QVBoxLayout(self)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(8)

        # -- numbering controls ---
        ctrl_group = QGroupBox("Capture Controls")
        ctrl_layout = QFormLayout(ctrl_group)

        self.next_num_edit = QLineEdit()
        self.next_num_edit.setFont(QFont("Segoe UI", 13, QFont.Bold))
        self.next_num_edit.setAlignment(Qt.AlignCenter)
        self.next_num_edit.setStyleSheet(
            "background: #181825; color: #f38ba8; border: 2px solid #f38ba8; "
            "border-radius: 6px; padding: 6px;"
        )
        ctrl_layout.addRow("Next Req #:", self.next_num_edit)

        self.sub_check = QCheckBox("Sub-requirement mode")
        ctrl_layout.addRow(self.sub_check)

        self.sub_parent_label = QLabel("")
        self.sub_parent_label.setFont(QFont("Segoe UI", 9))
        ctrl_layout.addRow(self.sub_parent_label)

        # -- markup color selector --
        color_row = QHBoxLayout()
        self._color_group = QButtonGroup(self)
        self._color_group.setExclusive(True)
        self._markup_colors = MARKUP_COLORS
        for i, (name, _rgb, hexc) in enumerate(MARKUP_COLORS):
            btn = QPushButton()
            btn.setFixedSize(28, 28)
            btn.setCheckable(True)
            btn.setToolTip(name)
            btn.setStyleSheet(
                f"QPushButton {{ background: {hexc}; border: 2px solid #45475a; "
                f"border-radius: 14px; }}"
                f"QPushButton:checked {{ border: 3px solid #cdd6f4; }}"
                f"QPushButton:hover {{ border-color: #89b4fa; }}"
            )
            if i == 0:
                btn.setChecked(True)
            self._color_group.addButton(btn, i)
            color_row.addWidget(btn)
        color_row.addStretch()
        ctrl_layout.addRow("Markup Color:", color_row)
        self.selected_markup_color = MARKUP_COLORS[0][1]
        self._color_group.buttonClicked[int].connect(self._on_markup_color)

        self._markup_hex = MARKUP_COLORS[0][2]

        layout.addWidget(ctrl_group)

        # -- requirements list ---
        list_label = QLabel("Captured Requirements")
        list_label.setFont(QFont("Segoe UI", 10, QFont.Bold))
        layout.addWidget(list_label)

        self.list_widget = QListWidget()
        self.list_widget.setSelectionMode(QListWidget.ExtendedSelection)
        self.list_widget.setSpacing(3)
        self.list_widget.setContextMenuPolicy(Qt.CustomContextMenu)
        self.list_widget.customContextMenuRequested.connect(
            self._show_context_menu
        )
        layout.addWidget(self.list_widget, 1)

        self.delete_btn = QPushButton("Delete Selected")
        self.delete_btn.setStyleSheet(
            "QPushButton { background: #45273a; color: #f38ba8; "
            "border: 1px solid #f38ba850; border-radius: 6px; padding: 6px 16px; }"
            "QPushButton:hover { background: #5a2e48; border-color: #f38ba8; }"
            "QPushButton:pressed { background: #6b354f; }"
        )
        self.delete_btn.clicked.connect(self._on_delete)
        layout.addWidget(self.delete_btn)

        self.list_widget.itemDoubleClicked.connect(self._on_double_click)

    def refresh(self, requirements: List[Requirement]):
        self.list_widget.clear()
        for req in requirements:
            item_widget = ReqItemWidget(req)
            item = QListWidgetItem()
            item.setSizeHint(item_widget.sizeHint())
            self.list_widget.addItem(item)
            self.list_widget.setItemWidget(item, item_widget)

    def _on_markup_color(self, btn_id):
        _name, rgb, hexc = self._markup_colors[btn_id]
        self.selected_markup_color = rgb
        self._markup_hex = hexc
        self.next_num_edit.setStyleSheet(
            f"background: #181825; color: {hexc}; border: 2px solid {hexc}; "
            "border-radius: 6px; padding: 6px;"
        )

    def _on_delete(self):
        rows = sorted(
            set(self.list_widget.row(item)
                for item in self.list_widget.selectedItems())
        )
        if rows:
            self.bulk_delete_requested.emit(rows)

    def _on_double_click(self, item):
        row = self.list_widget.row(item)
        if row >= 0:
            self.edit_requested.emit(row)

    def _show_context_menu(self, pos):
        item = self.list_widget.itemAt(pos)
        if item is None:
            return
        row = self.list_widget.row(item)
        if row < 0:
            return

        menu = QMenu(self)
        act_notes = menu.addAction("Edit Notes...")
        act_edit = menu.addAction("Edit Screenshot...")
        act_copy = menu.addAction("Copy Text")

        color_menu = menu.addMenu("Change Color")
        color_actions = []
        for i, (name, _rgb, _hexc) in enumerate(MARKUP_COLORS):
            act = color_menu.addAction(name)
            act.setData(i)
            color_actions.append(act)

        menu.addSeparator()
        act_delete = menu.addAction("Delete")

        action = menu.exec_(self.list_widget.mapToGlobal(pos))
        if action is None:
            return
        if action == act_notes:
            self.edit_notes_requested.emit(row)
        elif action == act_edit:
            self.edit_requested.emit(row)
        elif action == act_copy:
            self.copy_text_requested.emit(row)
        elif action == act_delete:
            self.delete_requested.emit(row)
        elif action in color_actions:
            self.change_color_requested.emit(row, action.data())


# ---------------------------------------------------------------------------
# Main Window
# ---------------------------------------------------------------------------

class MainWindow(QMainWindow):
    SCREENSHOT_ZOOM = 3.0  # render zoom for high-res screenshots

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Requirements Tracker")
        self.resize(1400, 900)

        # -- state --
        self._pdf_path: Optional[str] = None
        self._original_bytes: Optional[bytes] = None
        self._doc: Optional[fitz.Document] = None
        self._markup_path: Optional[str] = None
        self._requirements: List[Requirement] = []
        self._next_main = 1
        self._next_sub = 1
        self._last_main = 0
        self._unsaved_changes = False
        self._deleted_stack: list = []  # undo stack for deleted requirements

        # -- widgets --
        self._build_toolbar()
        self._build_central()
        self._build_statusbar()
        self._connect_signals()
        self._update_number_display()
        self._update_title()

        # drag-and-drop PDF opening
        self.setAcceptDrops(True)

    # ===================== UI construction =================================

    def _build_toolbar(self):
        tb = self.addToolBar("Main")
        tb.setMovable(False)
        tb.setIconSize(QSize(20, 20))
        style = self.style()

        # file actions
        self._act_open = tb.addAction(
            style.standardIcon(QStyle.SP_DialogOpenButton), "Open PDF"
        )
        self._act_save = tb.addAction(
            style.standardIcon(QStyle.SP_DialogSaveButton), "Save Markup"
        )
        self._act_export = tb.addAction(
            style.standardIcon(QStyle.SP_FileDialogDetailedView),
            "Export Requirements Doc"
        )
        tb.addSeparator()

        # navigation
        self._act_prev = tb.addAction(
            style.standardIcon(QStyle.SP_ArrowLeft), "Prev Page"
        )
        self._page_label = QLabel(" Page - / - ")
        self._page_label.setFont(QFont("Segoe UI", 10))
        tb.addWidget(self._page_label)
        self._act_next = tb.addAction(
            style.standardIcon(QStyle.SP_ArrowRight), "Next Page"
        )
        tb.addSeparator()

        # zoom
        self._act_zoom_out = tb.addAction(
            style.standardIcon(QStyle.SP_ArrowDown), "Zoom Out"
        )
        self._zoom_label = QLabel(" 100% ")
        self._zoom_label.setFont(QFont("Segoe UI", 10))
        tb.addWidget(self._zoom_label)
        self._act_zoom_in = tb.addAction(
            style.standardIcon(QStyle.SP_ArrowUp), "Zoom In"
        )
        self._act_fit = tb.addAction("Fit Width")

    def _build_central(self):
        splitter = QSplitter(Qt.Horizontal)

        self._viewer = PDFViewer()
        splitter.addWidget(self._viewer)

        self._panel = RequirementsPanel()
        self._panel.setMinimumWidth(280)
        self._panel.setMaximumWidth(400)
        splitter.addWidget(self._panel)

        splitter.setStretchFactor(0, 3)
        splitter.setStretchFactor(1, 1)
        self.setCentralWidget(splitter)

    def _build_statusbar(self):
        self._status = self.statusBar()
        self._status.showMessage("Open a PDF to begin.")

    def _connect_signals(self):
        self._act_open.triggered.connect(self._open_pdf)
        self._act_save.triggered.connect(self._save_markup)
        self._act_export.triggered.connect(self._manual_export)
        self._act_prev.triggered.connect(self._viewer.prev_page)
        self._act_next.triggered.connect(self._viewer.next_page)
        self._act_zoom_in.triggered.connect(self._viewer.zoom_in)
        self._act_zoom_out.triggered.connect(self._viewer.zoom_out)
        self._act_fit.triggered.connect(self._viewer.fit_width)

        self._viewer.selection_made.connect(self._handle_selection)
        self._viewer.page_changed.connect(self._on_page_changed)
        self._panel.delete_requested.connect(self._delete_requirement)
        self._panel.sub_check.stateChanged.connect(
            lambda _: self._update_number_display()
        )
        self._panel.next_num_edit.editingFinished.connect(
            self._on_next_number_edited
        )
        self._panel.edit_requested.connect(self._edit_screenshot)
        self._panel.bulk_delete_requested.connect(self._delete_requirements)
        self._panel.edit_notes_requested.connect(self._edit_notes)
        self._panel.change_color_requested.connect(self._change_requirement_color)
        self._panel.copy_text_requested.connect(self._copy_requirement_text)
        self._panel.list_widget.currentRowChanged.connect(
            self._on_list_selection_changed
        )

    # ===================== File operations =================================

    def _open_pdf(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "Open PDF / Drawing / SOW", "",
            "PDF Files (*.pdf);;All Files (*)"
        )
        if not path:
            return
        self._open_pdf_from_path(path)

    def _open_pdf_from_path(self, path: str):
        try:
            with open(path, "rb") as f:
                self._original_bytes = f.read()
            doc = self._open_clean_doc(self._original_bytes)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to open PDF:\n{e}")
            return

        # reset state
        self._pdf_path = path
        self._markup_path = None
        self._doc = doc
        self._set_unsaved(False)
        self._requirements.clear()
        self._deleted_stack.clear()
        self._next_main = 1
        self._next_sub = 1
        self._last_main = 0
        self._panel.sub_check.setChecked(False)

        self._viewer.set_document(self._doc, page=0)
        self._viewer.fit_width()
        self._panel.refresh(self._requirements)
        self._update_number_display()
        self._status.showMessage(f"Opened: {os.path.basename(path)}")

        # check for session restore
        self._check_session_restore()

    # ===================== Selection / capture =============================

    def _handle_selection(self, page_num: int, pixmap_rect: QRectF):
        """Called when the user finishes drawing a rectangle."""
        if not self._doc:
            QMessageBox.information(
                self, "No PDF", "Open a PDF file first."
            )
            return

        zoom = self._viewer.render_zoom
        ox, oy = self._viewer.page_rect_origin

        # convert pixmap coords → PDF points (accounting for page CropBox origin)
        pdf_x0 = pixmap_rect.x() / zoom + ox
        pdf_y0 = pixmap_rect.y() / zoom + oy
        pdf_x1 = pixmap_rect.right() / zoom + ox
        pdf_y1 = pixmap_rect.bottom() / zoom + oy
        pdf_rect = (pdf_x0, pdf_y0, pdf_x1, pdf_y1)

        # capture a clean high-res screenshot from the ORIGINAL pdf
        screenshot = self._capture_clean(page_num, pdf_rect)
        if screenshot is None or screenshot.isNull():
            return

        # extract text from the captured region
        extracted = self._extract_text(page_num, pdf_rect, screenshot)

        # determine requirement number
        num_str = self._allocate_number()

        req = Requirement(
            number=num_str,
            screenshot=screenshot,
            page=page_num,
            pdf_rect=pdf_rect,
            text=extracted,
            markup_color=self._panel.selected_markup_color,
        )
        self._requirements.append(req)
        self._sort_requirements()

        # rebuild in-memory stamped view (no disk save)
        self._rebuild_view()
        self._panel.refresh(self._requirements)
        # select the newly captured requirement after sorting
        new_row = self._requirements.index(req)
        self._panel.list_widget.setCurrentRow(new_row)
        self._update_number_display()
        self._save_session()
        self._status.showMessage(
            f"Requirement {req.number} captured  (unsaved)"
        )

    def _capture_clean(self, page_num: int, pdf_rect: tuple) -> Optional[QPixmap]:
        """Render the given rectangle from the *original* PDF at high res."""
        try:
            doc = self._open_clean_doc(self._original_bytes)
            page = doc[page_num]
            clip = fitz.Rect(pdf_rect)
            mat = fitz.Matrix(self.SCREENSHOT_ZOOM, self.SCREENSHOT_ZOOM)
            pix = page.get_pixmap(matrix=mat, clip=clip, alpha=False)
            img = QImage(
                pix.samples, pix.width, pix.height,
                pix.stride, QImage.Format_RGB888
            ).copy()
            doc.close()
            return QPixmap.fromImage(img)
        except Exception as e:
            QMessageBox.warning(self, "Capture Error", str(e))
            return None

    def _extract_text(self, page_num: int, pdf_rect: tuple,
                      screenshot: QPixmap) -> str:
        """Extract text from the selected PDF region.

        Tries native PDF text extraction first. Falls back to OCR if the
        result is empty and pytesseract is available.
        """
        text = ""
        # 1) Native PDF text extraction
        try:
            doc = self._open_clean_doc(self._original_bytes)
            page = doc[page_num]
            text = page.get_text("text", clip=fitz.Rect(pdf_rect))
            doc.close()
        except Exception:
            pass

        if text.strip():
            return text.strip()

        # 2) OCR fallback via pytesseract
        if HAS_TESSERACT:
            try:
                bio = pixmap_to_bytes(screenshot)
                pil_img = Image.open(bio)
                text = pytesseract.image_to_string(pil_img)
            except Exception:
                pass

        return text.strip()

    # ===================== Numbering =======================================

    def _allocate_number(self) -> str:
        existing = {req.number for req in self._requirements}
        if self._panel.sub_check.isChecked() and self._last_main > 0:
            # skip sub-numbers already in use
            while f"{self._last_main}.{self._next_sub}" in existing:
                self._next_sub += 1
            num_str = f"{self._last_main}.{self._next_sub}"
            self._next_sub += 1
        else:
            # skip main numbers already in use
            while str(self._next_main) in existing:
                self._next_main += 1
            num_str = str(self._next_main)
            self._last_main = self._next_main
            self._next_main += 1
            self._next_sub = 1
        return num_str

    @staticmethod
    def _req_sort_key(req):
        """Sort key: '3' -> (3, 0), '3.1' -> (3, 1)."""
        parts = req.number.split(".", 1)
        try:
            main = int(parts[0])
            sub = int(parts[1]) if len(parts) > 1 else 0
        except ValueError:
            return (999999, 0)
        return (main, sub)

    def _sort_requirements(self):
        """Sort requirements by number."""
        self._requirements.sort(key=self._req_sort_key)

    def _renumber_requirements(self):
        """Sort then renumber all requirements sequentially, preserving main/sub structure."""
        self._sort_requirements()
        main_counter = 0
        sub_counter = 0
        current_main = 0
        for req in self._requirements:
            is_sub = "." in req.number
            if is_sub:
                sub_counter += 1
                req.number = f"{current_main}.{sub_counter}"
            else:
                main_counter += 1
                current_main = main_counter
                sub_counter = 0
                req.number = str(main_counter)
        # update internal state to follow on from the last assigned numbers
        self._next_main = main_counter + 1
        self._last_main = current_main
        self._next_sub = sub_counter + 1

    def _update_number_display(self):
        if self._panel.sub_check.isChecked() and self._last_main > 0:
            nxt = f"{self._last_main}.{self._next_sub}"
            self._panel.sub_parent_label.setText(
                f"Sub-requirements under {self._last_main}"
            )
        else:
            nxt = str(self._next_main)
            self._panel.sub_parent_label.setText("")
        self._panel.next_num_edit.setText(nxt)

    def _on_next_number_edited(self):
        """Handle user manually changing the next requirement number."""
        text = self._panel.next_num_edit.text().strip()
        if not text:
            self._update_number_display()
            return

        # Check for duplicates against existing requirement numbers
        existing = {req.number for req in self._requirements}
        if text in existing:
            self._status.showMessage(
                f"Number '{text}' is already in use"
            )
            self._update_number_display()
            return

        if "." in text:
            # sub-requirement format e.g. "3.2"
            parts = text.split(".", 1)
            try:
                main = int(parts[0])
                sub = int(parts[1])
            except ValueError:
                self._update_number_display()
                return
            self._last_main = main
            self._next_sub = sub
            if not self._panel.sub_check.isChecked():
                self._panel.sub_check.setChecked(True)
        else:
            try:
                num = int(text)
            except ValueError:
                self._update_number_display()
                return
            self._next_main = num
            self._next_sub = 1

    # ===================== PDF stamping / rebuild ==========================

    @staticmethod
    def _open_clean_doc(pdf_bytes: bytes):
        """Open a PDF from bytes with structure tree fully stripped.

        Removes StructTreeRoot, MarkInfo from the catalog and StructParents
        from every page, then round-trips through tobytes() so MuPDF never
        sees leftover structure tree references.
        """
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        try:
            needs_clean = False
            cat = doc.pdf_catalog()
            xref = doc.xref_get_key(cat, "StructTreeRoot")
            if xref[0] != "null":
                doc.xref_set_key(cat, "StructTreeRoot", "null")
                needs_clean = True
            xref_mark = doc.xref_get_key(cat, "MarkInfo")
            if xref_mark[0] != "null":
                doc.xref_set_key(cat, "MarkInfo", "null")
                needs_clean = True
            # remove per-page StructParents references
            for i in range(len(doc)):
                pxref = doc[i].xref
                sp = doc.xref_get_key(pxref, "StructParents")
                if sp[0] != "null":
                    doc.xref_set_key(pxref, "StructParents", "null")
                    needs_clean = True
            if needs_clean:
                # round-trip so MuPDF reparses a clean document
                clean_bytes = doc.tobytes(garbage=4, deflate=True, clean=True)
                doc.close()
                doc = fitz.open(stream=clean_bytes, filetype="pdf")
        except Exception:
            pass  # not a valid PDF catalog — proceed as-is
        return doc

    def _rebuild_view(self):
        """Recreate all stamps on a fresh copy and display (no disk save)."""
        try:
            doc = self._open_clean_doc(self._original_bytes)
            for req in self._requirements:
                page = doc[req.page]
                r = fitz.Rect(req.pdf_rect)
                if req.edited and req.highlight_pixmap is not None:
                    self._overlay_screenshot(page, r, req.highlight_pixmap)
                self._stamp_page(page, r, req.number, req.markup_color)

            cur = self._viewer.current_page
            if self._doc:
                self._doc.close()
            self._doc = doc
            self._viewer.set_document(self._doc, page=cur)
            self._set_unsaved(True)
        except Exception as e:
            QMessageBox.warning(self, "Rebuild Error", str(e))

    def _save_markup(self):
        """Save the marked-up PDF to disk (user-triggered)."""
        if not self._requirements:
            QMessageBox.information(
                self, "Nothing to save",
                "Capture some requirements first."
            )
            return

        # prompt for path if not yet chosen
        if not self._markup_path:
            base = os.path.splitext(self._pdf_path)[0]
            date_str = datetime.now().strftime("%Y.%m.%d")
            default = f"{base} {date_str}_markup.pdf"
            path, _ = QFileDialog.getSaveFileName(
                self, "Save Marked-Up PDF As", default,
                "PDF Files (*.pdf)"
            )
            if not path:
                return
            self._markup_path = path

        try:
            doc = self._open_clean_doc(self._original_bytes)
            for req in self._requirements:
                page = doc[req.page]
                r = fitz.Rect(req.pdf_rect)
                if req.edited and req.highlight_pixmap is not None:
                    self._overlay_screenshot(page, r, req.highlight_pixmap)
                self._stamp_page(page, r, req.number, req.markup_color)
            doc.save(self._markup_path)
            doc.close()
            self._set_unsaved(False)
            # remove session sidecar since work is now saved in the PDF
            sp = self._session_path()
            if sp and os.path.exists(sp):
                try:
                    os.unlink(sp)
                except OSError:
                    pass
            self._status.showMessage(
                f"Saved: {os.path.basename(self._markup_path)}"
            )
        except Exception as e:
            QMessageBox.warning(self, "Save Error", str(e))
            return

        # also export requirements doc alongside
        self._auto_export_docx()

    @staticmethod
    def _overlay_screenshot(page, sel_rect: fitz.Rect, pixmap: QPixmap):
        """Replace the selected region on the PDF page with the edited screenshot."""
        img_bytes = pixmap_to_bytes(pixmap)
        page.insert_image(sel_rect, stream=img_bytes.read())

    @staticmethod
    def _stamp_page(page, sel_rect: fitz.Rect, number: str,
                    color: tuple = (0.85, 0.15, 0.15)):
        """Draw a dashed outline and a numbered stamp on a PDF page."""
        white = (1, 1, 1)

        # dashed outline around captured area
        page.draw_rect(sel_rect, color=color, width=0.75, dashes="[3 3]")

        # stamp label
        fontsize = 10
        fontname = "helv"
        text_w = fitz.get_text_length(number, fontname=fontname, fontsize=fontsize)
        pad = 3

        # position stamp at upper-left corner of selection
        sx = sel_rect.x0
        sy = sel_rect.y0
        stamp = fitz.Rect(
            sx - 1,
            sy - fontsize - 2 * pad,
            sx + text_w + 2 * pad + 1,
            sy,
        )

        # clamp stamp so it doesn't go above page
        if stamp.y0 < 0:
            shift = -stamp.y0
            stamp.y0 += shift
            stamp.y1 += shift

        page.draw_rect(stamp, color=color, fill=white, width=1.5)
        # text baseline sits near the bottom of the stamp box
        page.insert_text(
            fitz.Point(stamp.x0 + pad, stamp.y1 - pad),
            number,
            fontsize=fontsize,
            fontname=fontname,
            color=color,
        )

    # ===================== Requirements document export ====================

    def _default_export_base(self) -> str:
        """Return default export basename: '{pdf_name} YYYY.MM.DD RQMT'."""
        if self._pdf_path:
            name = os.path.splitext(os.path.basename(self._pdf_path))[0]
            directory = os.path.dirname(self._pdf_path)
        elif self._markup_path:
            name = os.path.splitext(os.path.basename(self._markup_path))[0]
            directory = os.path.dirname(self._markup_path)
        else:
            name = "requirements"
            directory = ""
        date_str = datetime.now().strftime("%Y.%m.%d")
        base = os.path.join(directory, f"{name} RQMT {date_str}")
        return base

    def _auto_export_docx(self):
        if not HAS_DOCX or not self._markup_path:
            return
        path = self._default_export_base() + ".docx"
        self._export_docx(path)

    def _manual_export(self):
        if not self._requirements:
            QMessageBox.information(
                self, "Nothing to export",
                "Capture some requirements first."
            )
            return

        # build file-type filters based on available libraries
        filters = []
        if HAS_DOCX:
            filters.append("Word Documents (*.docx)")
        if HAS_OPENPYXL:
            filters.append("Excel Workbooks (*.xlsx)")
        if not filters:
            QMessageBox.warning(
                self, "Missing Dependency",
                "Install python-docx or openpyxl to export:\n"
                "  pip install python-docx openpyxl"
            )
            return

        base = self._default_export_base()
        default_ext = ".docx" if HAS_DOCX else ".xlsx"
        path, chosen_filter = QFileDialog.getSaveFileName(
            self, "Export Requirements Document", base + default_ext,
            ";;".join(filters),
        )
        if not path:
            return

        if path.lower().endswith(".xlsx"):
            self._export_xlsx(path)
        else:
            self._export_docx(path)
        self._status.showMessage(f"Exported: {os.path.basename(path)}")

    def _export_docx(self, path: str):
        try:
            doc = DocxDocument()

            # title
            doc.add_heading("Requirements Tracker", level=0)
            if self._pdf_path:
                doc.add_paragraph(
                    f"Source: {os.path.basename(self._pdf_path)}"
                )
            doc.add_paragraph(
                f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            )
            doc.add_paragraph(
                f"Total requirements: {len(self._requirements)}"
            )
            doc.add_paragraph("")

            # table
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            headers = ["Req #", "Screenshot", "Extracted Text", "Page", "Notes"]
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for req in self._requirements:
                row = table.add_row()
                row.cells[0].text = req.number

                # embed screenshot
                img_io = pixmap_to_bytes(req.screenshot)
                paragraph = row.cells[1].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(img_io, width=Inches(3.5))

                row.cells[2].text = req.text
                row.cells[3].text = str(req.page + 1)
                row.cells[4].text = req.notes

            doc.save(path)
        except Exception as e:
            QMessageBox.warning(
                self, "Export Error", f"Failed to export:\n{e}"
            )

    def _export_xlsx(self, path: str):
        try:
            import tempfile
            wb = Workbook()
            ws = wb.active
            ws.title = "Requirements"

            # header info rows
            ws.append(["Requirements Tracker"])
            ws["A1"].font = XlFont(size=16, bold=True)
            if self._pdf_path:
                ws.append([f"Source: {os.path.basename(self._pdf_path)}"])
            ws.append([f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}"])
            ws.append([f"Total requirements: {len(self._requirements)}"])
            ws.append([])  # blank row

            # table header
            header_row = ws.max_row + 1
            headers = ["Req #", "Screenshot", "Extracted Text", "Page", "Notes"]
            ws.append(headers)
            for col_idx, h in enumerate(headers, 1):
                cell = ws.cell(row=header_row, column=col_idx)
                cell.font = XlFont(bold=True)

            # column widths
            ws.column_dimensions["A"].width = 10
            ws.column_dimensions["B"].width = 50
            ws.column_dimensions["C"].width = 50
            ws.column_dimensions["D"].width = 8
            ws.column_dimensions["E"].width = 30

            # data rows
            tmp_files = []
            for req in self._requirements:
                data_row = ws.max_row + 1
                ws.cell(row=data_row, column=1, value=req.number)

                # save screenshot to temp file for openpyxl
                img_io = pixmap_to_bytes(req.screenshot)
                tmp = tempfile.NamedTemporaryFile(
                    suffix=".png", delete=False
                )
                tmp.write(img_io.read())
                tmp.close()
                tmp_files.append(tmp.name)

                img = XlImage(tmp.name)
                # scale to fit ~350px wide, keep aspect ratio
                max_w = 350
                scale = min(max_w / img.width, 1.0)
                img.width = int(img.width * scale)
                img.height = int(img.height * scale)
                ws.add_image(img, f"B{data_row}")

                # set row height to fit the image
                ws.row_dimensions[data_row].height = img.height * 0.75

                ws.cell(row=data_row, column=3, value=req.text)
                ws.cell(row=data_row, column=3).alignment = XlAlignment(
                    wrap_text=True, vertical="top"
                )
                ws.cell(row=data_row, column=4, value=req.page + 1)
                ws.cell(row=data_row, column=5, value=req.notes)

            wb.save(path)

            # clean up temp files
            for f in tmp_files:
                try:
                    os.unlink(f)
                except OSError:
                    pass
        except Exception as e:
            QMessageBox.warning(
                self, "Export Error", f"Failed to export Excel:\n{e}"
            )

    # ===================== Delete / Undo ====================================

    def _delete_requirement(self, row: int):
        if 0 <= row < len(self._requirements):
            removed = self._requirements.pop(row)
            self._deleted_stack.append([(row, removed)])
            self._renumber_requirements()
            self._rebuild_view()
            self._panel.refresh(self._requirements)
            self._update_number_display()
            self._save_session()
            self._status.showMessage(
                f"Deleted requirement {removed.number}  (Ctrl+Z to undo)"
            )

    def _delete_requirements(self, rows: list):
        """Bulk delete multiple selected requirements (undo-capable)."""
        if not rows:
            return
        group = []
        for r in sorted(rows, reverse=True):
            if 0 <= r < len(self._requirements):
                removed = self._requirements.pop(r)
                group.append((r, removed))
        if group:
            group.reverse()  # store in ascending order for undo
            self._deleted_stack.append(group)
            self._renumber_requirements()
            self._rebuild_view()
            self._panel.refresh(self._requirements)
            self._update_number_display()
            self._save_session()
            self._status.showMessage(
                f"Deleted {len(group)} requirement(s)  (Ctrl+Z to undo)"
            )

    def _undo_delete(self):
        """Restore the last deleted requirement(s)."""
        if not self._deleted_stack:
            self._status.showMessage("Nothing to undo")
            return
        group = self._deleted_stack.pop()
        for orig_row, req in group:
            idx = min(orig_row, len(self._requirements))
            self._requirements.insert(idx, req)
        self._renumber_requirements()
        self._rebuild_view()
        self._panel.refresh(self._requirements)
        self._update_number_display()
        self._save_session()
        nums = ", ".join(req.number for _, req in group)
        self._status.showMessage(f"Restored requirement(s): {nums}")

    # ===================== Screenshot editing ================================

    def _edit_screenshot(self, row: int):
        if not (0 <= row < len(self._requirements)):
            return
        req = self._requirements[row]
        dlg = ScreenshotEditorDialog(req.screenshot, self)
        if dlg.exec_() == QDialog.Accepted:
            req.screenshot = dlg.get_pixmap()
            req.highlight_pixmap = dlg.get_highlight_pixmap()
            req.edited = True
            self._rebuild_view()
            self._panel.refresh(self._requirements)
            self._panel.list_widget.setCurrentRow(row)
            self._set_unsaved(True)
            self._save_session()
            self._status.showMessage(
                f"Requirement {req.number} screenshot edited  (unsaved)"
            )

    # ===================== Notes / Color / Copy =============================

    def _edit_notes(self, row: int):
        if not (0 <= row < len(self._requirements)):
            return
        req = self._requirements[row]
        dlg = QDialog(self)
        dlg.setWindowTitle(f"Notes — Requirement {req.number}")
        dlg.resize(400, 250)
        layout = QVBoxLayout(dlg)
        text_edit = QTextEdit()
        text_edit.setPlainText(req.notes)
        layout.addWidget(text_edit)
        btn_layout = QHBoxLayout()
        btn_ok = QPushButton("OK")
        btn_cancel = QPushButton("Cancel")
        btn_layout.addStretch()
        btn_layout.addWidget(btn_ok)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)
        btn_ok.clicked.connect(dlg.accept)
        btn_cancel.clicked.connect(dlg.reject)
        if dlg.exec_() == QDialog.Accepted:
            req.notes = text_edit.toPlainText()
            self._panel.refresh(self._requirements)
            self._panel.list_widget.setCurrentRow(row)
            self._set_unsaved(True)
            self._save_session()
            self._status.showMessage(
                f"Notes updated for requirement {req.number}"
            )

    def _change_requirement_color(self, row: int, color_idx: int):
        if not (0 <= row < len(self._requirements)):
            return
        if not (0 <= color_idx < len(MARKUP_COLORS)):
            return
        req = self._requirements[row]
        req.markup_color = MARKUP_COLORS[color_idx][1]
        self._rebuild_view()
        self._panel.refresh(self._requirements)
        self._panel.list_widget.setCurrentRow(row)
        self._save_session()
        self._status.showMessage(
            f"Color changed for requirement {req.number}"
        )

    def _copy_requirement_text(self, row: int):
        if not (0 <= row < len(self._requirements)):
            return
        req = self._requirements[row]
        clipboard = QApplication.clipboard()
        clipboard.setText(req.text)
        self._status.showMessage(
            f"Text copied for requirement {req.number}"
        )

    # ===================== Navigation / UI updates =========================

    def _set_unsaved(self, value: bool):
        self._unsaved_changes = value
        self._update_title()

    def _update_title(self):
        title = "Requirements Tracker"
        if self._pdf_path:
            title = f"{os.path.basename(self._pdf_path)} — {title}"
        if self._unsaved_changes:
            title = f"* {title}"
        self.setWindowTitle(title)

    def _go_to_page_dialog(self):
        if not self._doc:
            return
        total = len(self._doc)
        page, ok = QInputDialog.getInt(
            self, "Go to Page", f"Page (1–{total}):",
            self._viewer.current_page + 1, 1, total
        )
        if ok:
            self._viewer.go_to_page(page - 1)

    def _on_page_changed(self, current: int, total: int):
        self._page_label.setText(f" Page {current + 1} / {total} ")
        zoom_pct = int(self._viewer.render_zoom / PDFViewer.RENDER_ZOOM * 100)
        self._zoom_label.setText(f" {zoom_pct}% ")

    def _on_list_selection_changed(self, row: int):
        if 0 <= row < len(self._requirements):
            req = self._requirements[row]
            self._viewer.go_to_page(req.page)
            self._viewer.scroll_to_pdf_point(req.pdf_rect[0], req.pdf_rect[1])

    # ===================== Drag-and-drop ===================================

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                if url.toLocalFile().lower().endswith(".pdf"):
                    event.acceptProposedAction()
                    return
        event.ignore()

    def dragMoveEvent(self, event):
        event.acceptProposedAction()

    def dropEvent(self, event):
        for url in event.mimeData().urls():
            path = url.toLocalFile()
            if path.lower().endswith(".pdf"):
                self._open_pdf_from_path(path)
                return

    # ===================== Session save / restore ==========================

    def _session_path(self) -> Optional[str]:
        if self._pdf_path:
            return self._pdf_path + ".rqmt.json"
        return None

    @staticmethod
    def _pixmap_to_b64(pixmap: QPixmap) -> str:
        ba = QByteArray()
        buf = QBuffer(ba)
        buf.open(QBuffer.WriteOnly)
        pixmap.save(buf, "PNG")
        buf.close()
        return base64.b64encode(bytes(ba)).decode("ascii")

    @staticmethod
    def _b64_to_pixmap(b64_str: str) -> QPixmap:
        data = base64.b64decode(b64_str)
        img = QImage()
        img.loadFromData(data, "PNG")
        return QPixmap.fromImage(img)

    def _save_session(self):
        sp = self._session_path()
        if not sp:
            return
        try:
            data = []
            for req in self._requirements:
                entry = {
                    "number": req.number,
                    "page": req.page,
                    "pdf_rect": list(req.pdf_rect),
                    "text": req.text,
                    "edited": req.edited,
                    "markup_color": list(req.markup_color),
                    "notes": req.notes,
                    "screenshot": self._pixmap_to_b64(req.screenshot),
                }
                if req.highlight_pixmap is not None:
                    entry["highlight_pixmap"] = self._pixmap_to_b64(
                        req.highlight_pixmap
                    )
                data.append(entry)
            with open(sp, "w", encoding="utf-8") as f:
                json.dump(data, f)
        except Exception:
            pass  # non-critical; don't disrupt the user

    def _load_session(self) -> List[Requirement]:
        sp = self._session_path()
        if not sp or not os.path.exists(sp):
            return []
        try:
            with open(sp, "r", encoding="utf-8") as f:
                data = json.load(f)
            reqs = []
            for entry in data:
                screenshot = self._b64_to_pixmap(entry["screenshot"])
                highlight = None
                if "highlight_pixmap" in entry:
                    highlight = self._b64_to_pixmap(entry["highlight_pixmap"])
                reqs.append(Requirement(
                    number=entry["number"],
                    screenshot=screenshot,
                    page=entry["page"],
                    pdf_rect=tuple(entry["pdf_rect"]),
                    text=entry.get("text", ""),
                    edited=entry.get("edited", False),
                    highlight_pixmap=highlight,
                    markup_color=tuple(entry.get("markup_color",
                                                 (0.85, 0.15, 0.15))),
                    notes=entry.get("notes", ""),
                ))
            return reqs
        except Exception:
            return []

    def _check_session_restore(self):
        sp = self._session_path()
        if not sp or not os.path.exists(sp):
            return
        reply = QMessageBox.question(
            self, "Restore Session",
            "A previous session was found for this PDF.\n"
            "Would you like to restore your captured requirements?",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes,
        )
        if reply == QMessageBox.Yes:
            reqs = self._load_session()
            if reqs:
                self._requirements = reqs
                self._sort_requirements()
                # set numbering state from loaded requirements
                max_main = 0
                max_sub = {}
                for req in self._requirements:
                    parts = req.number.split(".", 1)
                    try:
                        m = int(parts[0])
                    except ValueError:
                        continue
                    max_main = max(max_main, m)
                    if len(parts) > 1:
                        try:
                            s = int(parts[1])
                        except ValueError:
                            continue
                        max_sub[m] = max(max_sub.get(m, 0), s)
                self._next_main = max_main + 1
                self._last_main = max_main
                self._next_sub = max_sub.get(max_main, 0) + 1
                self._rebuild_view()
                self._panel.refresh(self._requirements)
                self._update_number_display()
                self._set_unsaved(True)
                self._status.showMessage(
                    f"Restored {len(reqs)} requirement(s) from session"
                )
        else:
            # user declined; delete the stale sidecar
            try:
                os.unlink(sp)
            except OSError:
                pass

    # ===================== Keyboard shortcuts ==============================

    def keyPressEvent(self, event):
        key = event.key()
        mod = event.modifiers()

        if mod == Qt.ControlModifier:
            if key == Qt.Key_O:
                self._open_pdf()
                return
            if key == Qt.Key_S:
                self._save_markup()
                return
            if key == Qt.Key_E:
                self._manual_export()
                return
            if key == Qt.Key_Z:
                self._undo_delete()
                return
            if key == Qt.Key_G:
                self._go_to_page_dialog()
                return
            if key == Qt.Key_Equal or key == Qt.Key_Plus:
                self._viewer.zoom_in()
                return
            if key == Qt.Key_Minus:
                self._viewer.zoom_out()
                return

        if key == Qt.Key_PageDown or key == Qt.Key_Right:
            self._viewer.next_page()
            return
        if key == Qt.Key_PageUp or key == Qt.Key_Left:
            self._viewer.prev_page()
            return
        if key == Qt.Key_Delete:
            selected = self._panel.list_widget.selectedItems()
            if len(selected) > 1:
                rows = [self._panel.list_widget.row(item)
                        for item in selected]
                self._delete_requirements(rows)
            else:
                row = self._panel.list_widget.currentRow()
                if row >= 0:
                    self._delete_requirement(row)
            return
        if key == Qt.Key_F:
            self._viewer.fit_width()
            return

        super().keyPressEvent(event)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")

    # dark theme stylesheet
    app.setStyleSheet("""
        * {
            font-family: "Segoe UI", "Inter", sans-serif;
        }

        QMainWindow {
            background: #181825;
        }

        QToolBar {
            background: #1e1e2e;
            border-bottom: 1px solid #313244;
            spacing: 8px;
            padding: 6px 8px;
        }
        QToolBar QLabel {
            color: #cdd6f4;
        }
        QToolBar::separator {
            width: 1px;
            background: #313244;
            margin: 4px 6px;
        }

        QToolBar QToolButton {
            color: #cdd6f4;
            background: transparent;
            border: 1px solid transparent;
            border-radius: 6px;
            padding: 5px 10px;
        }
        QToolBar QToolButton:hover {
            background: #313244;
            border-color: #45475a;
        }
        QToolBar QToolButton:pressed {
            background: #45475a;
        }

        QGroupBox {
            font-weight: bold;
            color: #cdd6f4;
            background: #1e1e2e;
            border: 1px solid #313244;
            border-radius: 8px;
            margin-top: 12px;
            padding-top: 16px;
        }
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 12px;
            padding: 0 6px;
            color: #a6adc8;
        }

        QLabel {
            color: #cdd6f4;
        }

        QLineEdit {
            background: #1e1e2e;
            color: #cdd6f4;
            border: 1px solid #313244;
            border-radius: 6px;
            padding: 6px 8px;
            selection-background-color: #45475a;
        }
        QLineEdit:focus {
            border-color: #89b4fa;
        }

        QCheckBox {
            color: #cdd6f4;
            spacing: 6px;
        }
        QCheckBox::indicator {
            width: 16px;
            height: 16px;
            border: 1px solid #45475a;
            border-radius: 4px;
            background: #1e1e2e;
        }
        QCheckBox::indicator:checked {
            background: #89b4fa;
            border-color: #89b4fa;
        }

        QListWidget {
            background: #181825;
            border: 1px solid #313244;
            border-radius: 8px;
            outline: none;
            padding: 4px;
        }
        QListWidget::item {
            background: #1e1e2e;
            border: 1px solid #313244;
            border-radius: 8px;
            margin: 2px 0px;
        }
        QListWidget::item:selected {
            background: #2a2a4a;
            border-color: #89b4fa;
        }
        QListWidget::item:hover {
            background: #232340;
            border-color: #45475a;
        }

        QPushButton {
            background: #313244;
            color: #cdd6f4;
            border: 1px solid #45475a;
            border-radius: 6px;
            padding: 6px 16px;
            font-weight: 500;
        }
        QPushButton:hover {
            background: #45475a;
            border-color: #585b70;
        }
        QPushButton:pressed {
            background: #585b70;
        }
        QPushButton:disabled {
            background: #1e1e2e;
            color: #585b70;
            border-color: #313244;
        }

        QStatusBar {
            background: #11111b;
            color: #a6adc8;
            border-top: 1px solid #313244;
            padding: 4px 8px;
        }

        QSplitter::handle {
            background: #313244;
            width: 2px;
        }
        QSplitter::handle:hover {
            background: #89b4fa;
        }

        QScrollArea {
            background: #181825;
            border: none;
        }

        QScrollBar:vertical {
            background: #181825;
            width: 10px;
            border: none;
            border-radius: 5px;
        }
        QScrollBar::handle:vertical {
            background: #313244;
            min-height: 30px;
            border-radius: 5px;
        }
        QScrollBar::handle:vertical:hover {
            background: #45475a;
        }
        QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
            height: 0px;
        }
        QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {
            background: none;
        }
        QScrollBar:horizontal {
            background: #181825;
            height: 10px;
            border: none;
            border-radius: 5px;
        }
        QScrollBar::handle:horizontal {
            background: #313244;
            min-width: 30px;
            border-radius: 5px;
        }
        QScrollBar::handle:horizontal:hover {
            background: #45475a;
        }
        QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
            width: 0px;
        }
        QScrollBar::add-page:horizontal, QScrollBar::sub-page:horizontal {
            background: none;
        }

        QMenu {
            background: #1e1e2e;
            color: #cdd6f4;
            border: 1px solid #313244;
            border-radius: 8px;
            padding: 4px;
        }
        QMenu::item {
            padding: 6px 24px 6px 12px;
            border-radius: 4px;
        }
        QMenu::item:selected {
            background: #313244;
        }
        QMenu::separator {
            height: 1px;
            background: #313244;
            margin: 4px 8px;
        }

        QDialog {
            background: #1e1e2e;
        }

        QTextEdit {
            background: #181825;
            color: #cdd6f4;
            border: 1px solid #313244;
            border-radius: 6px;
            padding: 6px;
            selection-background-color: #45475a;
        }
        QTextEdit:focus {
            border-color: #89b4fa;
        }

        QSlider::groove:horizontal {
            background: #313244;
            height: 6px;
            border-radius: 3px;
        }
        QSlider::handle:horizontal {
            background: #89b4fa;
            width: 16px;
            height: 16px;
            margin: -5px 0;
            border-radius: 8px;
        }
        QSlider::handle:horizontal:hover {
            background: #b4d0fb;
        }

        QToolTip {
            background: #313244;
            color: #cdd6f4;
            border: 1px solid #45475a;
            border-radius: 4px;
            padding: 4px 8px;
        }

        QMessageBox {
            background: #1e1e2e;
        }
        QMessageBox QLabel {
            color: #cdd6f4;
        }

        QInputDialog {
            background: #1e1e2e;
        }

        QSpinBox {
            background: #1e1e2e;
            color: #cdd6f4;
            border: 1px solid #313244;
            border-radius: 6px;
            padding: 4px 8px;
        }
        QSpinBox:focus {
            border-color: #89b4fa;
        }
    """)

    window = MainWindow()
    window.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
